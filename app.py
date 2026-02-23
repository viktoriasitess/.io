import streamlit as st
import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches
import io
import re
from PIL import Image

# Налаштування сторінки
st.set_page_config(page_title="Wattpad Downloader Pro", page_icon="📖")

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0 Safari/537.36",
}

def add_image_to_docx(img_tag, doc):
    url = img_tag.get('src') or img_tag.get('data-original')
    if not url: return
    try:
        if url.startswith("//"): url = "https:" + url
        res = requests.get(url, headers=HEADERS, timeout=10)
        if res.status_code == 200:
            img_data = io.BytesIO(res.content)
            img = Image.open(img_data)
            if img.mode != 'RGB': img = img.convert('RGB')
            
            # Стиснення до 500px для економії місця
            max_w = 500
            if img.width > max_w:
                ratio = max_w / img.width
                img = img.resize((max_w, int(img.height * ratio)), Image.Resampling.LANCZOS)
            
            out_io = io.BytesIO()
            img.save(out_io, format='JPEG', quality=50, optimize=True)
            out_io.seek(0)
            doc.add_picture(out_io, width=Inches(5.0))
    except: pass

def process_content(html_text, doc):
    soup = BeautifulSoup(html_text, "html.parser")
    for element in soup.find_all(['h2', 'p', 'img']):
        if element.name == 'h2':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'p':
            para = doc.add_paragraph(element.get_text())
        elif element.name == 'img':
            add_image_to_docx(element, doc)

# Інтерфейс сайту
st.title("📖 Wattpad Downloader & Compressor")
st.markdown("Вставте посилання на глави, щоб отримати стиснутий Word-файл.")

urls_input = st.text_area("Посилання (кожне з нового рядка):", height=150)

if st.button("Завантажити та обробити"):
    urls = [u.strip() for u in urls_input.splitlines() if u.strip()]
    if not urls:
        st.error("Будь ласка, додайте хоча б одне посилання.")
    else:
        doc = Document()
        progress_bar = st.progress(0)
        
        for idx, url in enumerate(urls):
            st.write(f"Обробка: {url}")
            try:
                # Отримання назви глави
                res = requests.get(url, headers=HEADERS, timeout=15)
                soup = BeautifulSoup(res.text, "html.parser")
                title = soup.find("h1").get_text(strip=True) if soup.find("h1") else f"Глава {idx+1}"
                
                # Отримання тексту через API
                m = re.search(r"/(\d+)-", url)
                if m:
                    chapter_id = m.group(1)
                    api_res = requests.get(f"https://www.wattpad.com/apiv2/storytext?id={chapter_id}", headers=HEADERS)
                    doc.add_heading(title, level=1)
                    process_content(api_res.text, doc)
            except Exception as e:
                st.error(f"Помилка в {url}: {e}")
            
            progress_bar.progress((idx + 1) / len(urls))
        
        # Збереження у файл
        target_file = io.BytesIO()
        doc.save(target_file)
        target_file.seek(0)
        
        st.success("Готово! Файл максимально стиснуто.")
        st.download_button(
            label="⬇️ Скачати Word (.docx)",
            data=target_file,
            file_name="wattpad_book.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )